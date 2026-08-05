"""Markdown → DOCX 定稿转换工具（初赛方案提交用）。

用 python-docx 将受控 Markdown 子集（标题/表格/列表/引用/代码块/粗体/行内代码/段落）
转为 Word 文档，无 pandoc 依赖。设计参照 `convert_docx_to_md.py` 但方向相反。

用法：
    python scripts/md_to_docx.py docs/initial-round-proposal.md -o docs/initial-round-proposal.docx
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor


def _add_rich_text(paragraph, text: str) -> None:
    """解析行内标记（**粗体**、`行内代码`）写入段落 run。"""
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    for token in pattern.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(token)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """将 markdown 表格写入 docx 表格（首行为表头）。"""
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            _add_rich_text(p, row[j] if j < len(row) else "")


def convert(md_text: str) -> Document:
    """解析 markdown 文本并返回生成的 docx Document。"""
    doc = Document()
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        # 代码块
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过闭合 ```
            for code_line in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(code_line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            continue
        # 表格块：当前行与下一行都是 | 开头，且下一行是分隔符
        if (
            stripped.startswith("|")
            and i + 1 < n
            and re.match(r"^\s*\|[\s:\-|]+\|\s*$", lines[i + 1])
        ):
            rows: list[list[str]] = []
            # 表头行
            rows.append([c.strip() for c in stripped.strip("|").split("|")])
            i += 2  # 跳过分隔行
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            _add_table(doc, rows)
            continue
        # 标题
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            p = doc.add_heading(level=level)
            _add_rich_text(p, heading.group(2))
            i += 1
            continue
        # 引用
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(stripped.lstrip("> ").strip())
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            i += 1
            continue
        # 列表
        if stripped:
            list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
            if list_match:
                indent = len(list_match.group(1))
                text = list_match.group(3)
                style = (
                    "List Bullet"
                    if list_match.group(2) in ("-", "*") else "List Number"
                )
                p = doc.add_paragraph(style=style)
                p.paragraph_format.left_indent = Pt(18 + indent * 14)
                _add_rich_text(p, text)
                i += 1
                continue
        # 分隔线
        if stripped in ("---", "***"):
            doc.add_paragraph().add_run("―" * 20)
            i += 1
            continue
        # 空行
        if not stripped:
            i += 1
            continue
        # 普通段落
        p = doc.add_paragraph()
        _add_rich_text(p, stripped)
        i += 1
    return doc


def main() -> None:
    parser = argparse.ArgumentParser(description="Markdown → DOCX 定稿转换")
    parser.add_argument("input", type=Path, help="输入 markdown 文件")
    parser.add_argument("-o", "--output", type=Path, help="输出 docx 路径（默认同目录同名 .docx）")
    args = parser.parse_args()

    src = args.input
    if not src.exists():
        raise SystemExit(f"输入文件不存在: {src}")
    out = args.output or src.with_suffix(".docx")

    doc = convert(src.read_text(encoding="utf-8"))
    doc.save(out)
    print(f"已生成: {out}（段落数约 {len(doc.paragraphs)}）")


if __name__ == "__main__":
    main()
