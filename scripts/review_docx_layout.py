"""初赛方案 docx 排版审阅工具。

对 `md_to_docx.py` 生成的 Word 定稿做可复现的排版核验，输出 Markdown 审阅报告。
核验维度：
1. 文档结构：标题层级顺序（H1→H2→H3 无跳级）、章节数量
2. 表格：行数 / 列数 / 表头空值 / 单元格溢出风险（列数过多或超长文本）
3. 残留标记：markdown 语法残留（`**`、`|`、`` ` ``、`#`、`>`、列表符）
4. 字符卫生：emoji / 控制字符 / 空段落 / 超长段落
5. 对照源 markdown：章节标题一一对应，docx 内容不缺失

用法：
    python scripts/review_docx_layout.py docs/initial-round-proposal.docx \
        --source docs/initial-round-proposal.md -o docs/docx-layout-review.md
"""
from __future__ import annotations

import argparse
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

# 需要检查的 markdown 残留模式（docx 中不应出现）
MD_RESIDUE_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("粗体标记", re.compile(r"\*\*")),
    ("表格分隔符", re.compile(r"^\s*\|[\s:\-|]+\|\s*$", re.MULTILINE)),
    ("行内代码", re.compile(r"`")),
    ("标题标记", re.compile(r"^#{1,6}\s", re.MULTILINE)),
    ("引用标记", re.compile(r"^>\s?", re.MULTILINE)),
    ("列表标记", re.compile(r"^(\s*)([-*]|\d+\.)\s+", re.MULTILINE)),
]

# emoji / 装饰符号 / 控制字符
EMOJI_RE = re.compile(r"[\U0001F300-\U0001FAFF\u2600-\u27BF\uFE0F]")
CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")

# 超长阈值（中文字符宽度估算）
LONG_PARA_THRESHOLD = 200
LONG_CELL_THRESHOLD = 60
WIDE_TABLE_COLS = 6


def _estimate_chars(text: str) -> int:
    """估算显示宽度：中文/全角按 2，其余按 1。"""
    width = 0
    for ch in text:
        width += 2 if ord(ch) > 0x2E7F else 1
    return width


def review_docx(docx_path: Path) -> dict:
    """核验 docx 排版，返回结构化报告。"""
    doc = Document(docx_path)
    report: dict = {
        "file": str(docx_path),
        "n_paragraphs": len(doc.paragraphs),
        "n_tables": len(doc.tables),
        "headings": [],
        "issues": [],
        "notes": [],
        "tables": [],
    }

    # 1. 标题层级顺序检查
    prev_level = 0
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text or not style.startswith("Heading"):
            continue
        level = int(style.split()[-1]) if style.split()[-1].isdigit() else 1
        if level > prev_level + 1 and prev_level > 0:
            report["issues"].append(
                f"标题跳级：'{text[:40]}' 为 H{level}，上一标题为 H{prev_level}"
            )
        report["headings"].append({"level": level, "text": text})
        prev_level = level

    # 2. 表格检查
    for i, table in enumerate(doc.tables):
        n_rows, n_cols = len(table.rows), len(table.columns)
        header_empty = sum(
            1 for c in table.rows[0].cells if not c.text.strip()
        ) if n_rows else 0
        long_cells = 0
        for row in table.rows:
            for cell in row.cells:
                if _estimate_chars(cell.text) > LONG_CELL_THRESHOLD:
                    long_cells += 1
        table_info = {
            "idx": i,
            "rows": n_rows,
            "cols": n_cols,
            "header_empty": header_empty,
            "long_cells": long_cells,
        }
        report["tables"].append(table_info)
        if n_cols >= WIDE_TABLE_COLS:
            report["issues"].append(
                f"表格 {i} 列数过多（{n_cols} 列），Word 排版易溢出，建议拆表或转列表"
            )
        if header_empty:
            report["issues"].append(f"表格 {i} 表头存在 {header_empty} 个空单元格")
        if long_cells:
            report["issues"].append(
                f"表格 {i} 有 {long_cells} 个超长单元格（>60 显示宽度），可能换行过多"
            )

    # 3. 残留标记 + 字符卫生（跳过标题样式段落，避免「1. 问题」标题被列表正则误报）
    paragraphs_text = [p.text for p in doc.paragraphs]
    para_styles = [
        (p.style.name if p.style else "") or "" for p in doc.paragraphs
    ]
    for name, pattern in MD_RESIDUE_PATTERNS:
        for text, style in zip(paragraphs_text, para_styles):
            if not text or style.startswith("Heading"):
                continue
            if pattern.search(text):
                report["issues"].append(
                    f"markdown 残留（{name}）：{text[:60]!r}"
                )
    for i, p in enumerate(paragraphs_text):
        if EMOJI_RE.search(p):
            report["issues"].append(f"段落 {i} 含 emoji/装饰符号：{p[:60]!r}")
        if CONTROL_RE.search(p):
            report["issues"].append(f"段落 {i} 含控制字符：{p[:60]!r}")
    n_empty = sum(1 for p in paragraphs_text if not p.strip())
    if n_empty:
        report["issues"].append(f"存在 {n_empty} 个空段落")
    for p in paragraphs_text:
        if _estimate_chars(p) > LONG_PARA_THRESHOLD:
            report["notes"].append(
                f"长段落提示（{_estimate_chars(p)} 显示宽度，Word 自动换行可接受）：{p[:60]}..."
            )

    return report


def compare_with_source(report: dict, source_path: Path) -> dict:
    """对照源 markdown 章节，检查 docx 标题是否完整覆盖。"""
    src = source_path.read_text(encoding="utf-8")
    src_headings = [
        line.lstrip("# ").strip()
        for line in src.splitlines()
        if re.match(r"^#{1,4}\s", line.strip())
    ]
    docx_headings = [h["text"] for h in report["headings"]]
    missing = [
        h for h in src_headings
        if not any(h.split("（")[0][:12] in d or d[:12] in h[:12] for d in docx_headings)
    ]
    report["source"] = str(source_path)
    report["src_headings"] = src_headings
    report["missing_in_docx"] = missing
    if missing:
        report["issues"].append(f"docx 缺失源 markdown 标题：{missing}")
    return report


def render_markdown(report: dict) -> str:
    """渲染审阅报告 Markdown。"""
    lines = [
        "# 初赛方案 docx 排版审阅报告",
        "",
        f"> 生成时间：{datetime.now(timezone.utc).isoformat()}",
        f"> 审阅对象：`{report['file']}`",
        "",
        "## 1. 文档结构",
        "",
        f"- 段落数：{report['n_paragraphs']}",
        f"- 表格数：{report['n_tables']}",
        f"- 标题数：{len(report['headings'])}",
        "",
        "### 标题层级",
        "",
    ]
    for h in report["headings"]:
        lines.append(f"- H{h['level']}：{h['text'][:60]}")
    lines += ["", "## 2. 表格核验", ""]
    if report["tables"]:
        lines.append("| 表格 | 行数 | 列数 | 表头空值 | 超长单元格 |")
        lines.append("|------|------|------|---------|-----------|")
        for t in report["tables"]:
            lines.append(
                f"| {t['idx']} | {t['rows']} | {t['cols']} | "
                f"{t['header_empty']} | {t['long_cells']} |"
            )
    else:
        lines.append("（无表格）")
    lines += ["", "## 3. 问题清单", ""]
    if report["issues"]:
        for i, issue in enumerate(report["issues"], 1):
            lines.append(f"{i}. {issue}")
    else:
        lines.append("未发现排版问题。")
    if report.get("notes"):
        lines += ["", "### 提示（非问题，Word 自动换行可接受）", ""]
        for i, note in enumerate(report["notes"], 1):
            lines.append(f"{i}. {note}")
    lines += ["", "## 4. 结论", ""]
    n_issues = len(report["issues"])
    if n_issues == 0:
        verdict = "✅ 排版合格，可提交"
    elif n_issues <= 3:
        verdict = "🟡 轻微问题，建议修复后提交"
    else:
        verdict = "🔴 存在较多排版问题，需修复后重新生成"
    lines += [f"- 问题数：{n_issues}", f"- **结论：{verdict}**", ""]
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description="初赛 docx 排版审阅")
    parser.add_argument("input", type=Path, help="docx 文件路径")
    parser.add_argument("--source", type=Path, help="源 markdown（可选，对照章节）")
    parser.add_argument("-o", "--output", type=Path, help="审阅报告输出路径")
    args = parser.parse_args()

    if not args.input.exists():
        raise SystemExit(f"docx 不存在: {args.input}")
    report = review_docx(args.input)
    if args.source and args.source.exists():
        report = compare_with_source(report, args.source)
    md = render_markdown(report)
    out = args.output or args.input.with_suffix(".layout-review.md")
    out.write_text(md, encoding="utf-8")
    print(md)
    print(f"\n审阅报告已生成: {out}")


if __name__ == "__main__":
    main()
