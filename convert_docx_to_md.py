#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert DOCX file to Markdown using python-docx
"""

import os
import sys

from docx import Document
from docx.oxml.ns import qn


def get_run_style(run):
    """Check if run has bold, italic, underline formatting"""
    styles = []
    if run.bold:
        styles.append('bold')
    if run.italic:
        styles.append('italic')
    if run.underline:
        styles.append('underline')
    return styles


def process_run(run):
    """Process a single run, applying markdown formatting"""
    text = run.text
    if not text:
        return ''

    styles = get_run_style(run)

    result = text

    # Apply formatting from inside out: underline -> italic -> bold
    if 'bold' in styles:
        result = f'**{result}**'
    if 'italic' in styles:
        result = f'*{result}*'
    if 'underline' in styles:
        result = f'<u>{result}</u>'

    return result


def process_paragraph(para, list_info=None):
    """Process a paragraph and return markdown text"""
    style_name = para.style.name if para.style else 'Normal'
    text_parts = []

    # Process hyperlinks and runs
    for child in para._element.iterchildren():
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'hyperlink':
            # Get hyperlink URL
            rid = child.get(qn('r:id'))
            url = ''
            if rid and para.part.rels:
                rel = para.part.rels.get(rid)
                if rel:
                    url = rel.target_ref

            # Process runs inside hyperlink
            link_text = ''
            for run in child.iter(qn('w:r')):
                run_obj = para.runs[0]  # dummy
                run_obj._r = run
                link_text += process_run(run_obj)

            if url and link_text:
                text_parts.append(f'[{link_text}]({url})')
        elif tag == 'r':
            # Find the corresponding run object
            for run in para.runs:
                if run._r is child:
                    text_parts.append(process_run(run))
                    break

    text = ''.join(text_parts) if text_parts else para.text

    # Handle list items
    if list_info:
        is_ordered, level = list_info
        indent = '  ' * level
        if is_ordered:
            return f'{indent}1. {text}'
        else:
            return f'{indent}- {text}'

    # Handle headings
    if style_name.startswith('Heading'):
        try:
            level = int(style_name.replace('Heading', '').strip())
            level = max(1, min(6, level))
            return '#' * level + ' ' + text
        except ValueError:
            pass
    elif style_name == 'Title':
        return '# ' + text
    elif style_name == 'Subtitle':
        return '## ' + text

    return text


def is_paragraph_list_item(para):
    """Check if paragraph is a list item and return (is_ordered, level) or None"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return None

    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None

    ilvl = numPr.find(qn('w:ilvl'))
    numId = numPr.find(qn('w:numId'))

    if ilvl is None or numId is None:
        return None

    level = int(ilvl.get(qn('w:val'), '0'))
    num_id_val = numId.get(qn('w:val'), '0')

    # Determine if ordered by checking the numbering XML
    try:
        numbering_part = para.part.document.part.numbering_part
        if numbering_part:
            num = numbering_part._numbering.find(qn('w:num'), {qn('w:numId'): num_id_val})
            if num is not None:
                abstract_num_id_elem = num.find(qn('w:abstractNumId'))
                if abstract_num_id_elem is not None:
                    abstract_num_id = abstract_num_id_elem.get(qn('w:val'))
                    abstract_num = numbering_part._numbering.find(
                        qn('w:abstractNum'), {qn('w:abstractNumId'): abstract_num_id}
                    )
                    if abstract_num is not None:
                        lvl_elem = abstract_num.find(
                            qn('w:lvl'), {qn('w:ilvl'): str(level)}
                        )
                        if lvl_elem is not None:
                            numFmt = lvl_elem.find(qn('w:numFmt'))
                            if numFmt is not None:
                                fmt_val = numFmt.get(qn('w:val'))
                                is_ordered = fmt_val not in ('bullet', 'none')
                                return (is_ordered, level)
    except Exception:
        pass

    return (False, level)


def process_table(table):
    """Process a table and return markdown table text"""
    rows = []
    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # Get cell text, replacing newlines with <br>
            cell_text = cell.text.strip().replace('\n', '<br>')
            cells.append(cell_text)
        rows.append(cells)

    if not rows:
        return ''

    # Build markdown table
    md_lines = []
    # Header row
    md_lines.append('| ' + ' | '.join(rows[0]) + ' |')
    # Separator
    md_lines.append('| ' + ' | '.join(['---'] * len(rows[0])) + ' |')
    # Data rows
    for row in rows[1:]:
        md_lines.append('| ' + ' | '.join(row) + ' |')

    return '\n'.join(md_lines)


def convert_docx_to_markdown(docx_path, md_path):
    """Convert docx file to markdown file"""
    doc = Document(docx_path)
    md_content = []

    # Track whether we're in a list to add blank lines around it
    in_list = False

    for element in doc.element.body.iterchildren():
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # Find the corresponding paragraph object
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break

            if para is None:
                continue

            list_info = is_paragraph_list_item(para)
            text = process_paragraph(para, list_info)

            if list_info:
                in_list = True
                md_content.append(text)
            else:
                if in_list:
                    # End of list, add blank line
                    md_content.append('')
                    in_list = False
                if text.strip():
                    md_content.append(text)
                    md_content.append('')  # blank line after paragraph
                else:
                    md_content.append('')  # preserve empty lines

        elif tag == 'tbl':
            # End list if in one
            if in_list:
                md_content.append('')
                in_list = False

            # Find corresponding table
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break

            if table:
                md_content.append(process_table(table))
                md_content.append('')  # blank line after table

    result = '\n'.join(md_content)

    # Clean up excessive blank lines
    import re
    result = re.sub(r'\n{4,}', '\n\n\n', result)
    result = result.strip() + '\n'

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(result)

    return result


def main():
    if len(sys.argv) < 2:
        print('Usage: python convert_docx_to_md.py <input.docx> [output.md]')
        sys.exit(1)

    input_path = sys.argv[1]
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        base = os.path.splitext(input_path)[0]
        output_path = base + '.md'

    print(f'Converting: {input_path} -> {output_path}')
    result = convert_docx_to_markdown(input_path, output_path)
    print(f'Done! Output size: {len(result)} characters')
    print(f'Output file: {output_path}')


if __name__ == '__main__':
    main()
